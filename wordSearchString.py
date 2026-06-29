#!/usr/bin/env python3

import argparse
import sys
from pathlib import Path

from docx import Document


def extract_text_from_docx(file_path: Path):
    """
    Yield text blocks from a .docx file.
    Includes paragraphs and table cell content.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"ERROR: Could not read {file_path}: {e}", file=sys.stderr)
        return

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            yield text

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    yield text


def find_docx_files(paths, recursive=False):
    """
    Yield .docx files from the provided input paths.
    """
    for input_path in paths:
        path = Path(input_path)

        if not path.exists():
            print(f"WARNING: Path does not exist: {path}", file=sys.stderr)
            continue

        if path.is_file():
            if path.suffix.lower() == ".docx":
                yield path
            else:
                print(f"WARNING: Skipping non-.docx file: {path}", file=sys.stderr)

        elif path.is_dir():
            if recursive:
                iterator = path.rglob("*.docx")
            else:
                iterator = path.glob("*.docx")

            for file_path in iterator:
                if file_path.is_file():
                    yield file_path


def search_in_file(file_path: Path, needle: str, ignore_case=False):
    """
    Search for needle in a .docx file.
    Returns a list of matching text blocks.
    """
    matches = []

    if ignore_case:
        needle_cmp = needle.lower()
    else:
        needle_cmp = needle

    for block in extract_text_from_docx(file_path):
        haystack = block.lower() if ignore_case else block
        if needle_cmp in haystack:
            matches.append(block)

    return matches


def main():
    parser = argparse.ArgumentParser(
        description="Search for text in .docx files, similar to grep."
    )
    parser.add_argument(
        "search_string",
        help="The string to search for"
    )
    parser.add_argument(
        "paths",
        nargs="+",
        help="One or more .docx files and/or directories"
    )
    parser.add_argument(
        "-r", "--recursive",
        action="store_true",
        help="Recursively search directories"
    )
    parser.add_argument(
        "-i", "--ignore-case",
        action="store_true",
        help="Case-insensitive search"
    )
    parser.add_argument(
        "-l", "--files-with-matches",
        action="store_true",
        help="Only print filenames with matches"
    )

    args = parser.parse_args()

    found_any = False

    for file_path in find_docx_files(args.paths, recursive=args.recursive):
        matches = search_in_file(
            file_path=file_path,
            needle=args.search_string,
            ignore_case=args.ignore_case
        )

        if matches:
            found_any = True

            if args.files_with_matches:
                print(file_path)
            else:
                for match in matches:
                    print(f"{file_path}: {match}")

    sys.exit(0 if found_any else 1)


if __name__ == "__main__":
    main()
